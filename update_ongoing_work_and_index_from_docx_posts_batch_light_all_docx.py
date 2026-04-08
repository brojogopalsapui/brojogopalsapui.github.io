from pathlib import Path
import re

# ===== USER CONFIG =====
WEEKLY_INPUTS_DIR = Path("weekly-inputs")
POSTS_DIR = Path("posts")

# Batch mode: process every DOCX inside weekly-inputs/ except known templates
BATCH_PROCESS_ALL_DOCX = True

# Single-file fallback modes used only when batch mode is False
AUTO_PICK_LATEST_DOCX = False
MANUAL_DOCX_PATH = Path("weekly-inputs/2026-03-stack-safeguard-pipelines.docx")

INPUT_ONGOING_HTML = Path("ongoing-work.html")
INPUT_INDEX_HTML = Path("index.html")

# Best default: update the real website files in place
OUTPUT_ONGOING_HTML = INPUT_ONGOING_HTML
OUTPUT_INDEX_HTML = INPUT_INDEX_HTML

# Full consistency: always generate/update the matching posts/...html page
AUTO_GENERATE_FULL_POST = True
OUTPUT_POSTS_DIR = POSTS_DIR

# Keep True so rerunning the notebook updates the same note instead of duplicating it
REPLACE_DUPLICATE_POST_ID = True

# Homepage update controls
UPDATE_HOME_FLOATING = True
UPDATE_HOME_SLIDER = True
HOME_SLIDER_MAX_CARDS = 6

# Consistency controls
# If True, the DOCX filename stem becomes the default post ID:
# weekly-inputs/2026-03-stack-safeguard-pipelines.docx
# -> posts/2026-03-stack-safeguard-pipelines.html
FORCE_POST_ID_FROM_FILENAME = True
AUTO_USE_FILENAME_BASED_FULL_POST_LINK = True

TEMPLATE_DOCX_HINTS = {
    "weekly-research-watch-template.docx",
    "template.docx",
    "research-watch-template.docx",
}

def _docx_sort_key(path: Path):
    name = path.stem
    m = re.match(r"^(\d{4})-(\d{2})(?:-(.+))?$", name, re.IGNORECASE)
    if m:
        year = int(m.group(1))
        month = int(m.group(2))
        tail = (m.group(3) or "").lower()
        return (0, year, month, tail)
    return (1, path.stat().st_mtime, name.lower())

def list_processable_docx_files(folder: Path):
    candidates = []
    for f in folder.glob("*.docx"):
        if f.name.lower() in {name.lower() for name in TEMPLATE_DOCX_HINTS}:
            continue
        if f.name.startswith("~$"):
            continue
        candidates.append(f)
    if not candidates:
        raise FileNotFoundError(
            f"No usable DOCX files found in {folder}."
        )
    return sorted(candidates, key=_docx_sort_key)

def pick_latest_docx(folder: Path) -> Path:
    return list_processable_docx_files(folder)[-1]

if BATCH_PROCESS_ALL_DOCX:
    DOCX_PATHS = list_processable_docx_files(WEEKLY_INPUTS_DIR)
else:
    DOCX_PATHS = [pick_latest_docx(WEEKLY_INPUTS_DIR)] if AUTO_PICK_LATEST_DOCX else [MANUAL_DOCX_PATH]

LATEST_DOCX_PATH = DOCX_PATHS[-1]

print("DOCX files selected:")
for p in DOCX_PATHS:
    latest_flag = "   <-- homepage latest note" if p == LATEST_DOCX_PATH else ""
    print(f"- {p}{latest_flag}")


from docx import Document
from bs4 import BeautifulSoup, NavigableString
from html import escape
from calendar import month_name
import re

def _clean(value):
    if value is None:
        return ""
    if isinstance(value, str):
        return value.strip()
    return str(value).strip()

def stem_to_month_year(stem: str) -> str:
    m = re.match(r"^(\d{4})-(\d{2})-", stem)
    if not m:
        return stem
    year = int(m.group(1))
    month = int(m.group(2))
    if 1 <= month <= 12:
        return f"{month_name[month]} {year}"
    return stem

def parse_weekly_docx(path: Path):
    doc = Document(str(path))
    data = {"metadata": {}}

    if doc.tables:
        table = doc.tables[0]
        for row in table.rows[1:]:
            key = _clean(row.cells[0].text)
            val = _clean(row.cells[1].text)
            if key:
                data["metadata"][key] = val

    wanted_headings = [
        "Preview",
        "Full Note Paragraph 1",
        "Full Note Paragraph 2",
        "What Is Changing Technically",
        "What Reviewers Should Notice",
        "Current Research Tension",
    ]

    paragraphs = doc.paragraphs
    for i, p in enumerate(paragraphs):
        heading = _clean(p.text)
        if heading in wanted_headings:
            j = i + 1
            content = []
            while j < len(paragraphs):
                t = _clean(paragraphs[j].text)
                if t in wanted_headings or t in {"Metadata", "How the notebook uses this DOCX"}:
                    break
                if t:
                    content.append(t)
                j += 1

            if heading in {"What Is Changing Technically", "What Reviewers Should Notice"}:
                data[heading] = content
            else:
                data[heading] = "\n".join(content)

    return data

def normalize_docx_data(data, docx_stem: str):
    md = data["metadata"]
    warnings = []

    if FORCE_POST_ID_FROM_FILENAME:
        existing_post_id = _clean(md.get("Post ID"))
        if existing_post_id and existing_post_id != docx_stem:
            warnings.append(
                f"Overriding DOCX Post ID '{existing_post_id}' with filename-based Post ID '{docx_stem}' for consistency."
            )
        md["Post ID"] = docx_stem
    else:
        if not _clean(md.get("Post ID")):
            md["Post ID"] = docx_stem

    if AUTO_USE_FILENAME_BASED_FULL_POST_LINK:
        existing_link = _clean(md.get("Full Post Link (optional)"))
        suggested = RELATIVE_FULL_POST_LINK
        if existing_link and existing_link != suggested:
            warnings.append(
                f"Overriding DOCX Full Post Link '{existing_link}' with filename-based link '{suggested}' for consistency."
            )
        md["Full Post Link (optional)"] = suggested

    if not _clean(md.get("Title")):
        md["Title"] = docx_stem.replace("-", " ").title()
        warnings.append("Title was blank in the DOCX, so a title was inferred from the filename.")

    if not _clean(md.get("Meta Line")):
        md["Meta Line"] = f"Research Watch • {stem_to_month_year(docx_stem)}"
        warnings.append("Meta Line was blank in the DOCX, so one was inferred from the filename.")

    if not _clean(data.get("Preview")) and _clean(md.get("Preview")):
        data["Preview"] = _clean(md.get("Preview"))

    return warnings

def _is_valid_href(href: str) -> bool:
    href = _clean(href)
    if not href:
        return False
    if href.startswith(("http://", "https://", "mailto:", "#")):
        return True
    if " " in href:
        return False
    return "/" in href or href.endswith(".html") or href.endswith(".pdf")

def add_link_html(links, href, label):
    href = _clean(href)
    label = _clean(label)
    if not _is_valid_href(href) or not label:
        return
    if href.startswith("http://") or href.startswith("https://"):
        links.append(
            f'<a href="{escape(href, quote=True)}" rel="noopener noreferrer" target="_blank">{escape(label)}</a>'
        )
    else:
        links.append(f'<a href="{escape(href, quote=True)}">{escape(label)}</a>')

def build_watch_article(data):
    md = data["metadata"]

    post_id = _clean(md.get("Post ID")) or DOCX_STEM
    meta_line = _clean(md.get("Meta Line")) or f"Research Watch • {stem_to_month_year(DOCX_STEM)}"
    category = _clean(md.get("Category"))
    if category:
        category_map = {"academic": "Academic Signals", "industry": "Industry & Innovation", "ecosystem": "Ecosystem & Releases"}
        category_label = category_map.get(category.lower(), category)
        if category_label.lower() not in meta_line.lower():
            meta_line = f"{meta_line} • {category_label}"
    title = _clean(md.get("Title")) or "Untitled research note"
    preview = _clean(data.get("Preview")) or _clean(md.get("Preview"))
    if not preview or preview == title:
        preview = (_clean(p1)[:220].rsplit(" ", 1)[0] + "...") if _clean(p1) else title
    p1 = _clean(data.get("Full Note Paragraph 1"))
    p2 = _clean(data.get("Full Note Paragraph 2"))
    tech_list = data.get("What Is Changing Technically", [])
    reviewer_list = data.get("What Reviewers Should Notice", [])
    tension = _clean(data.get("Current Research Tension"))

    links = []
    add_link_html(links, md.get("Full Post Link (optional)"), "Read full post")
    add_link_html(links, md.get("Related Static Page (optional)"), md.get("Related Static Page Label (optional)"))
    add_link_html(links, md.get("External Link 1 URL (optional)"), md.get("External Link 1 Label (optional)"))
    add_link_html(links, md.get("External Link 2 URL (optional)"), md.get("External Link 2 Label (optional)"))

    tech_items = "\n".join([f"                      <li>{escape(_clean(item))}</li>" for item in tech_list if _clean(item)])
    reviewer_items = "\n".join([f"                      <li>{escape(_clean(item))}</li>" for item in reviewer_list if _clean(item)])
    links_html = " ".join(links)

    article_html = f'''
<article class="watch-note accordion" id="{escape(post_id, quote=True)}">
  <button aria-expanded="false" class="accordion-trigger" type="button">
    <span class="accordion-meta">{escape(meta_line)}</span>
    <span class="accordion-title">{escape(title)}</span>
    <span class="accordion-preview">
      {escape(preview)}
    </span>
    <span class="accordion-cta">Read full note</span>
    <span aria-hidden="true" class="accordion-icon"></span>
  </button>

  <div aria-hidden="true" class="accordion-panel">
    <div class="accordion-panel-inner">
      <p>
        {escape(p1)}
      </p>

      <p>
        {escape(p2)}
      </p>

      <div class="watch-columns">
        <div class="watch-block">
          <h4>What is changing technically</h4>
          <ul>
{tech_items}
          </ul>
        </div>

        <div class="watch-block">
          <h4>What reviewers should notice</h4>
          <ul>
{reviewer_items}
          </ul>
        </div>
      </div>

      <div class="watch-bottom-note">
        <strong>Current research tension:</strong> {escape(tension)}
      </div>

      <div class="watch-inline-links">
        {links_html}
      </div>
    </div>
  </div>
</article>
'''.strip()

    return post_id, title, preview, article_html

def build_home_watch_card(post_id, title, preview):
    return f'''
<a class="post-card watch-card" href="ongoing-work.html#{escape(post_id, quote=True)}">
  <span class="tag">Research Watch</span>
  <h3>{escape(title)}</h3>
  <p>
    {escape(preview)}
  </p>
</a>
'''.strip()

def build_full_post_html(data):
    md = data["metadata"]

    post_id = _clean(md.get("Post ID")) or DOCX_STEM
    title = _clean(md.get("Title")) or "Untitled research note"
    preview = _clean(data.get("Preview")) or _clean(md.get("Preview"))
    if not preview or preview == title:
        preview = (_clean(p1)[:220].rsplit(" ", 1)[0] + "...") if _clean(p1) else title
    p1 = _clean(data.get("Full Note Paragraph 1"))
    p2 = _clean(data.get("Full Note Paragraph 2"))
    tech_list = data.get("What Is Changing Technically", [])
    reviewer_list = data.get("What Reviewers Should Notice", [])
    tension = _clean(data.get("Current Research Tension"))
    month_year = stem_to_month_year(DOCX_STEM)

    related_links = []
    add_link_html(related_links, md.get("Related Static Page (optional)"), md.get("Related Static Page Label (optional)"))
    add_link_html(related_links, md.get("External Link 1 URL (optional)"), md.get("External Link 1 Label (optional)"))
    add_link_html(related_links, md.get("External Link 2 URL (optional)"), md.get("External Link 2 Label (optional)"))
    related_links_html = " ".join(related_links) if related_links else ""
    related_links_block = (
        '<div class="btn-row" style="margin-top:1rem;">' + related_links_html + "</div>"
        if related_links_html else ""
    )

    tech_items = "\n".join([f"              <li>{escape(_clean(item))}</li>" for item in tech_list if _clean(item)])
    reviewer_items = "\n".join([f"              <li>{escape(_clean(item))}</li>" for item in reviewer_list if _clean(item)])

    html = f'''<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1.0" />
  <title>{escape(title)} | Brojogopal Sapui</title>
  <meta name="description" content="{escape(preview)}" />
  <link rel="preconnect" href="https://fonts.googleapis.com" />
  <link rel="preconnect" href="https://fonts.gstatic.com" crossorigin />
  <link href="https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800&display=swap" rel="stylesheet" />
  <link rel="stylesheet" href="../assets/css/style.css" />
</head>
<body>
  <header class="site-header">
    <div class="container nav-wrap">
      <a class="brand" href="../index.html" aria-label="Brojogopal Sapui Home">B<span>S</span></a>
      <nav class="nav">
        <a href="../index.html">Home</a>
        <a href="../about.html">About</a>
        <a href="../research.html">Research</a>
        <a href="../ongoing-work.html">Trending Topics</a>
        <a href="../publications.html">Resources</a>
        <a href="../contact.html">Contact</a>
      </nav>
      <button class="menu-btn" aria-label="Toggle menu" aria-expanded="false">
        <span></span>
        <span></span>
        <span></span>
      </button>
    </div>
  </header>

  <main>
    <section class="page-hero">
      <div class="container">
        <span class="eyebrow">Research Watch • {escape(month_year)}</span>
        <h1>{escape(title)}</h1>
        <p class="lead">{escape(preview)}</p>
      </div>
    </section>

    <section class="section">
      <div class="container split">
        <div class="content-card">
          <span class="kicker">Overview</span>
          <h2>What is changing</h2>
          <p>{escape(p1)}</p>
          <p>{escape(p2)}</p>
        </div>

        <div class="content-card">
          <span class="kicker">Why it matters</span>
          <h2>Research significance</h2>
          <ul class="check-list">
{tech_items}
          </ul>
        </div>
      </div>
    </section>

    <section class="section alt">
      <div class="container">
        <div class="section-head">
          <span class="eyebrow">Discussion</span>
          <h2>What reviewers should notice</h2>
          <p>
            These review points help separate benchmark-level claims from stronger system-level conclusions.
          </p>
        </div>

        <div class="content-card">
          <ul class="check-list">
{reviewer_items}
          </ul>
        </div>
      </div>
    </section>

    <section class="section">
      <div class="container">
        <div class="content-card">
          <span class="kicker">Current research tension</span>
          <h2>Why this topic matters now</h2>
          <p>{escape(tension)}</p>
          {related_links_block}
        </div>
      </div>
    </section>

    <section class="section alt">
      <div class="container">
        <div class="cta">
          <div>
            <span class="eyebrow">Next Step</span>
            <h2>Back to ongoing research updates</h2>
            <p>
              Return to the running list of research-watch topics and evolving system-level notes.
            </p>
          </div>
          <div class="cta-actions">
            <a class="btn btn-primary" href="../ongoing-work.html#{escape(post_id, quote=True)}">Back to this note</a>
            <a class="btn btn-secondary" href="../research.html">Research</a>
          </div>
        </div>
      </div>
    </section>
  </main>

  <footer class="site-footer">
    <div class="container footer-grid">
      <div>
        <h3>Brojogopal Sapui</h3>
        <p>AI Security• Hardware Trust • Edge/Physical AI</p>
      </div>
      <div>
        <h4>Main Pages</h4>
        <a href="../research.html">Research</a>
        <a href="../ongoing-work.html">Trending Topics</a>
        <a href="../publications.html">Resources</a>
      </div>
      <div>
        <h4>Focus</h4>
        <p>Cross-layer AI security, trustworthy deployment, hardware-aware defense, and physical intelligence.</p>
      </div>
    </div>
  </footer>

  <script src="../assets/js/main.js"></script>
</body>
</html>'''.strip()

    return html

def detect_doctype(html_text: str) -> str:
    m = re.match(r'\s*(<!DOCTYPE[^>]+>)', html_text, flags=re.IGNORECASE)
    return m.group(1) if m else "<!DOCTYPE html>"

def validate_paths(docx_path, ongoing_input_path, ongoing_output_path, index_input_path, index_output_path):
    warnings = []

    if not docx_path.exists():
        raise FileNotFoundError(f"DOCX file not found: {docx_path.resolve()}")
    if not ongoing_input_path.exists():
        raise FileNotFoundError(f"Input ongoing-work.html not found: {ongoing_input_path.resolve()}")
    if not index_input_path.exists():
        raise FileNotFoundError(f"Input index.html not found: {index_input_path.resolve()}")

    if docx_path.name.lower() in {name.lower() for name in TEMPLATE_DOCX_HINTS}:
        warnings.append("DOCX_PATH looks like a blank template file. Use a filled weekly DOCX instead.")

    html_dir = ongoing_input_path.parent
    css_path = html_dir / "assets" / "css" / "style.css"
    js_path = html_dir / "assets" / "js" / "main.js"

    if not css_path.exists():
        warnings.append(f"CSS file not found at {css_path}. Local preview may look broken.")
    if not js_path.exists():
        warnings.append(f"JS file not found at {js_path}. Accordions/sliders may not work locally.")

    if ongoing_input_path.resolve() != ongoing_output_path.resolve():
        warnings.append("OUTPUT_ONGOING_HTML differs from INPUT_ONGOING_HTML. Best default: update in place.")
    if index_input_path.resolve() != index_output_path.resolve():
        warnings.append("OUTPUT_INDEX_HTML differs from INPUT_INDEX_HTML. Best default: update in place.")

    return warnings

def update_ongoing_work_html(input_html_path, output_html_path, article_id, article_html, replace_duplicate=True):
    html_text = input_html_path.read_text(encoding="utf-8")
    doctype = detect_doctype(html_text)
    soup = BeautifulSoup(html_text, "html.parser")

    watch_stack = soup.select_one("div.watch-stack")
    if watch_stack is None:
        raise ValueError("Could not find div.watch-stack in ongoing-work.html")

    if replace_duplicate:
        existing = watch_stack.find("article", {"id": article_id})
        if existing:
            existing.decompose()

    fragment = BeautifulSoup(article_html, "html.parser")
    new_article = fragment.find("article")
    first_existing = watch_stack.find("article")
    if first_existing:
        first_existing.insert_before("\n")
        first_existing.insert_before(new_article)
        first_existing.insert_before("\n\n          ")
    else:
        watch_stack.append(new_article)

    final_html = doctype + "\n" + str(soup)
    output_html_path.write_text(final_html, encoding="utf-8")
    return output_html_path

def update_index_html(input_html_path, output_html_path, post_id, title, preview,
                      update_floating=True, update_slider=True,
                      replace_duplicate=True, max_cards=6):
    html_text = input_html_path.read_text(encoding="utf-8")
    doctype = detect_doctype(html_text)
    soup = BeautifulSoup(html_text, "html.parser")

    if update_floating:
        # Support both the older floating-notif homepage layout and the newer
        # research-alert topbar layout without changing the site's theme or structure.
        notif_text = soup.select_one("div.floating-notif div.notif-text")
        if notif_text is not None:
            notif_text.clear()
            strong = soup.new_tag("strong")
            strong.string = "Research Watch:"
            notif_text.append(strong)
            notif_text.append(NavigableString(f" {title}. "))
            link = soup.new_tag("a", href=f"ongoing-work.html#{post_id}")
            link.string = "Read note →"
            notif_text.append(link)
        else:
            research_alert = soup.select_one("div.research-alert")
            if research_alert is None:
                raise ValueError(
                    "Could not find a supported floating/latest-note block in index.html "
                    "(expected either .floating-notif .notif-text or .research-alert)."
                )

            badge = research_alert.select_one(".research-alert-badge")
            if badge is not None:
                badge.string = "Latest note"

            chips = research_alert.select_one(".hero-chips")
            if chips is None:
                chips = soup.new_tag("div", attrs={"class": "hero-chips"})
                link_node = research_alert.select_one("a.research-alert-link")
                if link_node is not None:
                    link_node.insert_before("\n")
                    link_node.insert_before(chips)
                    link_node.insert_before("\n")
                else:
                    research_alert.append(chips)

            chips.clear()
            chip_span = soup.new_tag("span")
            chip_span.string = preview or title
            chips.append(chip_span)

            link_node = research_alert.select_one("a.research-alert-link")
            if link_node is None:
                link_node = soup.new_tag("a", attrs={"class": "research-alert-link"})
                research_alert.append("\n")
                research_alert.append(link_node)
            link_node["href"] = f"ongoing-work.html#{post_id}"
            link_node.string = "Read the research watch →"

    if update_slider:
        watch_track = soup.select_one("div#watchTrack")
        if watch_track is not None:
            if replace_duplicate:
                for a in watch_track.select("a.watch-card"):
                    href = a.get("href", "")
                    if href.endswith(f"#{post_id}"):
                        a.decompose()

            card_fragment = BeautifulSoup(build_home_watch_card(post_id, title, preview), "html.parser")
            new_card = card_fragment.find("a")
            first_card = watch_track.find("a", class_="watch-card")
            if first_card:
                first_card.insert_before("\n")
                first_card.insert_before(new_card)
                first_card.insert_before("\n")
            else:
                watch_track.append(new_card)

            cards = watch_track.select("a.watch-card")
            if max_cards and len(cards) > max_cards:
                for card in cards[max_cards:]:
                    card.decompose()
        else:
            # Newer homepage versions may not include a watch slider. Skip this
            # step quietly so the notebook still works.
            pass

    final_html = doctype + "\n" + str(soup)
    output_html_path.write_text(final_html, encoding="utf-8")
    return output_html_path

def write_full_post_html(output_path: Path, full_post_html: str):
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(full_post_html, encoding="utf-8")
    return output_path

path_warnings = validate_paths(
    LATEST_DOCX_PATH,
    INPUT_ONGOING_HTML,
    OUTPUT_ONGOING_HTML,
    INPUT_INDEX_HTML,
    OUTPUT_INDEX_HTML,
)

if path_warnings:
    print("CONFIG WARNINGS:")
    for w in path_warnings:
        print("-", w)
    print()

processed_items = []

for docx_path in DOCX_PATHS:
    CURRENT_DOCX_STEM = docx_path.stem
    RELATIVE_FULL_POST_LINK = f"posts/{CURRENT_DOCX_STEM}.html"
    suggested_post_html_path = OUTPUT_POSTS_DIR / f"{CURRENT_DOCX_STEM}.html"

    data = parse_weekly_docx(docx_path)
    normalize_warnings = normalize_docx_data(data, CURRENT_DOCX_STEM)

    md = data["metadata"]
    if normalize_warnings:
        print(f"NORMALIZATION NOTES for {docx_path.name}:")
        for w in normalize_warnings:
            print("-", w)
        print()

    processed_items.append(
        {
            "docx_path": docx_path,
            "docx_stem": CURRENT_DOCX_STEM,
            "suggested_post_html_path": suggested_post_html_path,
            "data": data,
            "post_id": _clean(md.get("Post ID")),
            "title": _clean(md.get("Title")),
            "meta_line": _clean(md.get("Meta Line")),
            "preview": _clean(data.get("Preview")) or _clean(md.get("Preview")),
            "full_post_link": _clean(md.get("Full Post Link (optional)")),
        }
    )

print("Selected DOCX summary:")
for item in processed_items:
    latest_tag = " [homepage latest note]" if item["docx_path"] == LATEST_DOCX_PATH else ""
    print(f"- {item['docx_path'].name}{latest_tag}")
    print("  Post ID:", item["post_id"])
    print("  Title:", item["title"])
    print("  Meta Line:", item["meta_line"])
    print("  Full post link:", item["full_post_link"])
    print("  Generated full post page:", item["suggested_post_html_path"])
    print()

ongoing_output_path = None
index_output_path = None
latest_home_item = None
generated_post_paths = []

for item in processed_items:
    CURRENT_DOCX_STEM = item["docx_stem"]
    DOCX_STEM = CURRENT_DOCX_STEM
    RELATIVE_FULL_POST_LINK = f"posts/{CURRENT_DOCX_STEM}.html"

    post_id, title, preview, article_html = build_watch_article(item["data"])

    ongoing_input_for_this_run = OUTPUT_ONGOING_HTML if ongoing_output_path is not None else INPUT_ONGOING_HTML
    ongoing_output_path = update_ongoing_work_html(
        input_html_path=ongoing_input_for_this_run,
        output_html_path=OUTPUT_ONGOING_HTML,
        article_id=post_id,
        article_html=article_html,
        replace_duplicate=REPLACE_DUPLICATE_POST_ID,
    )

    index_input_for_this_run = OUTPUT_INDEX_HTML if index_output_path is not None else INPUT_INDEX_HTML
    index_output_path = update_index_html(
        input_html_path=index_input_for_this_run,
        output_html_path=OUTPUT_INDEX_HTML,
        post_id=post_id,
        title=title,
        preview=preview,
        update_floating=False,
        update_slider=UPDATE_HOME_SLIDER,
        replace_duplicate=REPLACE_DUPLICATE_POST_ID,
        max_cards=HOME_SLIDER_MAX_CARDS,
    )

    if AUTO_GENERATE_FULL_POST:
        full_post_html = build_full_post_html(item["data"])
        generated_post_path = write_full_post_html(item["suggested_post_html_path"], full_post_html)
        generated_post_paths.append(generated_post_path)

    if item["docx_path"] == LATEST_DOCX_PATH:
        latest_home_item = {
            "post_id": post_id,
            "title": title,
            "preview": preview,
        }

if latest_home_item is None:
    raise RuntimeError("Could not determine which processed DOCX should control the homepage latest note.")

index_input_for_floating = OUTPUT_INDEX_HTML if index_output_path is not None else INPUT_INDEX_HTML
index_output_path = update_index_html(
    input_html_path=index_input_for_floating,
    output_html_path=OUTPUT_INDEX_HTML,
    post_id=latest_home_item["post_id"],
    title=latest_home_item["title"],
    preview=latest_home_item["preview"],
    update_floating=UPDATE_HOME_FLOATING,
    update_slider=False,
    replace_duplicate=REPLACE_DUPLICATE_POST_ID,
    max_cards=HOME_SLIDER_MAX_CARDS,
)

print("Updated ongoing-work.html:", ongoing_output_path.resolve())
print("Updated index.html:", index_output_path.resolve())
if generated_post_paths:
    print("Generated/updated full posts:")
    for p in generated_post_paths:
        print("-", p.resolve())
print()
print("Homepage latest note now points to:", latest_home_item["post_id"])
print("Homepage latest note title:", latest_home_item["title"])
print()
print("Next step: open the updated files from the repo root, review them, then push to GitHub.")